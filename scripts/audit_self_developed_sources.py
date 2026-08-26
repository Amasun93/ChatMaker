#!/usr/bin/env python3
"""Verify that the curated 23-module evidence still matches the raw handoff files."""

from __future__ import annotations

import argparse
from collections import Counter
import hashlib
import json
from pathlib import Path
import sys

from docx import Document
from pypdf import PdfReader
import yaml


ROOT = Path(__file__).resolve().parents[1]
MANIFEST = ROOT / "knowledge_sources/catalogs/self-developed-hardware.yaml"


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def inspect_document(path: Path) -> tuple[str, int]:
    suffix = path.suffix.casefold()
    if suffix == ".pdf":
        reader = PdfReader(path)
        if not reader.pages:
            raise ValueError("pdf_has_no_pages")
        extracted = "".join((page.extract_text() or "") for page in reader.pages[:3])
        return "pdf", len(extracted.strip())
    if suffix == ".docx":
        document = Document(path)
        extracted = "\n".join(paragraph.text for paragraph in document.paragraphs)
        extracted += "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        return "docx", len(extracted.strip())
    return suffix.lstrip(".") or "unknown", 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-root", type=Path, required=True)
    args = parser.parse_args(argv)
    source_root = args.source_root.resolve()
    manifest = yaml.safe_load(MANIFEST.read_text(encoding="utf-8"))
    unique: dict[str, dict[str, str]] = {}
    for module in manifest["modules"]:
        for evidence in module["source_evidence"]:
            unique.setdefault(evidence["path"], evidence)
    errors: list[str] = []
    types: Counter[str] = Counter()
    parsed: Counter[str] = Counter()
    text_bearing: Counter[str] = Counter()
    for relative, evidence in sorted(unique.items()):
        path = source_root / Path(relative)
        types[evidence["evidence_type"]] += 1
        if not path.is_file():
            errors.append(f"missing:{relative}")
            continue
        actual = digest(path)
        if actual != evidence["sha256"]:
            errors.append(f"sha256_mismatch:{relative}")
            continue
        try:
            kind, text_length = inspect_document(path)
        except Exception as exc:
            errors.append(f"unreadable:{relative}:{type(exc).__name__}:{exc}")
            continue
        if kind in {"pdf", "docx"}:
            parsed[kind] += 1
            if text_length:
                text_bearing[kind] += 1
    result = {
        "success": not errors,
        "module_count": manifest["module_count"],
        "unique_evidence_files": len(unique),
        "evidence_types": dict(sorted(types.items())),
        "parsed_documents": dict(sorted(parsed.items())),
        "text_bearing_documents": dict(sorted(text_bearing.items())),
        "errors": errors,
        "boundary": "Hash and readability checks do not prove wiring, firmware, or physical effects.",
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["success"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
